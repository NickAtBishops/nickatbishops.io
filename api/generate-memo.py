from http.server import BaseHTTPRequestHandler
import json
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            body = self.rfile.read(content_length)
            data = json.loads(body)
            
            docx_bytes = generate_memo(data)
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', 'attachment; filename="GGC_Investment_Memo.docx"')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(docx_bytes)
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()


def generate_memo(data):
    """Generate Investment Memo from JSON data"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    fin = data.get('financials', {})
    prop = fin.get('propertyInfo', {}) or data.get('propertyInfo', {})
    comps = data.get('comps', {})
    market = data.get('market', {})
    
    property_name = prop.get('name', 'Subject Property')
    
    # ============ TITLE PAGE ============
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('INVESTMENT MEMO')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)
    
    doc.add_paragraph()
    
    property_title = doc.add_paragraph()
    property_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = property_title.add_run(property_name)
    run.bold = True
    run.font.size = Pt(24)
    
    address = doc.add_paragraph()
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address.add_run(f"{prop.get('address', '')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = location.add_run(f"{prop.get('city', '')}, {prop.get('state', '')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    ggc = doc.add_paragraph()
    ggc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ggc.add_run('Gary Group Capital')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_page_break()
    
    # ============ EXECUTIVE SUMMARY ============
    add_heading(doc, 'Executive Summary')
    
    # Key metrics box
    income_total = sum(i.get('t12Amount', 0) or 0 for i in fin.get('income', []))
    expense_total = sum(e.get('t12Amount', 0) or 0 for e in fin.get('expenses', []))
    noi = income_total - expense_total
    asking = prop.get('askingPrice', 0) or data.get('propertyInfo', {}).get('askingPrice', 0)
    cap_rate = (noi / asking * 100) if asking else 0
    units = prop.get('totalUnits', 0) or data.get('propertyInfo', {}).get('totalUnits', 0)
    price_per_unit = asking / units if units else 0
    
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Asking Price', 'T12 NOI', 'Cap Rate', 'Price/Unit']
    values = [
        f"${asking:,.0f}",
        f"${noi:,.0f}",
        f"{cap_rate:.1f}%",
        f"${price_per_unit:,.0f}"
    ]
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D6E3F8")
    
    for i, value in enumerate(values):
        table.rows[1].cells[i].text = value
    
    doc.add_paragraph()
    
    # Recommendation
    rr = fin.get('rentRoll', {})
    occupancy = rr.get('occupancyRate', 0)
    demand = market.get('demandSignal', 'Unknown')
    
    rec_para = doc.add_paragraph()
    rec_para.add_run('Recommendation: ').bold = True
    
    if cap_rate > 7 and occupancy > 0.9:
        rec_para.add_run('PURSUE - Strong fundamentals with attractive cap rate')
    elif cap_rate > 6:
        rec_para.add_run('CONSIDER - Moderate opportunity, requires further diligence')
    else:
        rec_para.add_run('PASS - Returns do not meet investment criteria')
    
    doc.add_paragraph()
    
    # ============ PROPERTY OVERVIEW ============
    add_heading(doc, 'Property Overview')
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    overview_data = [
        ('Property Name', property_name),
        ('Address', f"{prop.get('address', '')}"),
        ('City, State', f"{prop.get('city', '')}, {prop.get('state', '')}"),
        ('Total Units/Sites', str(units)),
        ('Occupancy', f"{occupancy * 100 if occupancy < 1 else occupancy:.1f}%"),
        ('Avg Lot Rent', f"${rr.get('avgMonthlyRent', 0):,.0f}/mo")
    ]
    
    for i, (label, value) in enumerate(overview_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], "F2F2F2")
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # ============ FINANCIAL SUMMARY ============
    add_heading(doc, 'Financial Summary')
    
    doc.add_paragraph().add_run('Income').bold = True
    
    if fin.get('income'):
        table = doc.add_table(rows=len(fin['income']) + 2, cols=3)
        table.style = 'Table Grid'
        
        # Header
        for i, h in enumerate(['Category', 'T12 Amount', 'Notes']):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[0].cells[i], "D6E3F8")
        
        # Data
        for i, item in enumerate(fin['income'], 1):
            table.rows[i].cells[0].text = item.get('ggcCategory', item.get('sellerName', ''))
            table.rows[i].cells[1].text = f"${item.get('t12Amount', 0):,.0f}"
            table.rows[i].cells[2].text = item.get('notes', '')
        
        # Total
        table.rows[-1].cells[0].text = 'Total Income'
        table.rows[-1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[-1].cells[1].text = f"${income_total:,.0f}"
        table.rows[-1].cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Expenses').bold = True
    
    if fin.get('expenses'):
        table = doc.add_table(rows=len(fin['expenses']) + 2, cols=3)
        table.style = 'Table Grid'
        
        for i, h in enumerate(['Category', 'T12 Amount', 'Notes']):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[0].cells[i], "D6E3F8")
        
        for i, item in enumerate(fin['expenses'], 1):
            table.rows[i].cells[0].text = item.get('ggcCategory', item.get('sellerName', ''))
            table.rows[i].cells[1].text = f"${item.get('t12Amount', 0):,.0f}"
            table.rows[i].cells[2].text = item.get('notes', '')
        
        table.rows[-1].cells[0].text = 'Total Expenses'
        table.rows[-1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[-1].cells[1].text = f"${expense_total:,.0f}"
        table.rows[-1].cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    noi_para = doc.add_paragraph()
    noi_para.add_run('Net Operating Income: ').bold = True
    noi_para.add_run(f"${noi:,.0f}")
    
    doc.add_page_break()
    
    # ============ RENT COMPS ============
    add_heading(doc, 'Rent Comparables')
    
    rent_comps = comps.get('rentComps', [])
    if rent_comps:
        table = doc.add_table(rows=len(rent_comps) + 1, cols=5)
        table.style = 'Table Grid'
        
        for i, h in enumerate(['Property', 'Location', 'Units', 'Lot Rent', 'Occupancy']):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[0].cells[i], "D6E3F8")
        
        for i, comp in enumerate(rent_comps, 1):
            table.rows[i].cells[0].text = comp.get('name', '')
            table.rows[i].cells[1].text = f"{comp.get('city', '')}, {comp.get('state', '')}"
            table.rows[i].cells[2].text = str(comp.get('units', ''))
            table.rows[i].cells[3].text = f"${comp.get('lotRent', 0):,.0f}"
            occ = comp.get('occupancy', 0)
            table.rows[i].cells[4].text = f"{occ * 100 if occ < 1 else occ:.0f}%"
        
        doc.add_paragraph()
        conclusion = doc.add_paragraph()
        conclusion.add_run('Market Rent Conclusion: ').bold = True
        conclusion.add_run(f"${comps.get('marketRentConclusion', 0):,.0f}/month")
    else:
        doc.add_paragraph('No rent comparables available.')
    
    doc.add_paragraph()
    
    # ============ SALE COMPS ============
    add_heading(doc, 'Sale Comparables')
    
    sale_comps = comps.get('saleComps', [])
    if sale_comps:
        table = doc.add_table(rows=len(sale_comps) + 1, cols=5)
        table.style = 'Table Grid'
        
        for i, h in enumerate(['Property', 'Date', 'Units', 'Sale Price', 'Cap Rate']):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[0].cells[i], "D6E3F8")
        
        for i, comp in enumerate(sale_comps, 1):
            table.rows[i].cells[0].text = comp.get('name', '')
            table.rows[i].cells[1].text = comp.get('saleDate', '')
            table.rows[i].cells[2].text = str(comp.get('units', ''))
            table.rows[i].cells[3].text = f"${comp.get('salePrice', 0):,.0f}"
            cr = comp.get('capRate', 0)
            table.rows[i].cells[4].text = f"{cr * 100 if cr < 1 else cr:.1f}%"
        
        doc.add_paragraph()
        conclusion = doc.add_paragraph()
        conclusion.add_run('Market Cap Rate: ').bold = True
        mcr = comps.get('marketCapRateConclusion', 0)
        conclusion.add_run(f"{mcr * 100 if mcr < 1 else mcr:.1f}%")
    else:
        doc.add_paragraph('No sale comparables available.')
    
    doc.add_paragraph()
    
    # ============ MARKET ANALYSIS ============
    add_heading(doc, 'Market & Alternative Housing Analysis')
    
    demo = market.get('demographics', {})
    alt = market.get('altHousing', {})
    
    doc.add_paragraph().add_run('Demographics').bold = True
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    demo_data = [
        ('County Population', f"{demo.get('countyPopulation', 0):,}"),
        ('Population Growth (5yr)', f"{demo.get('populationGrowth', 0):.1f}%"),
        ('Median HH Income', f"${demo.get('medianHHIncome', 0):,}"),
        ('Unemployment Rate', f"{demo.get('unemploymentRate', 0):.1f}%"),
        ('Major Employers', ', '.join(demo.get('majorEmployers', [])[:3]))
    ]
    
    for i, (label, value) in enumerate(demo_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], "F2F2F2")
        table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Alternative Housing Comparison').bold = True
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    alt_data = [
        ('Avg Single Family Home', f"${alt.get('avgSFHomePrice', 0):,}"),
        ('Avg 1BR Apartment', f"${alt.get('avgRent1BR', 0):,}/mo"),
        ('Avg 2BR Apartment', f"${alt.get('avgRent2BR', 0):,}/mo"),
        ('MHP All-In Cost', f"${alt.get('mhpAllInCost', 0):,}/mo")
    ]
    
    for i, (label, value) in enumerate(alt_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], "F2F2F2")
        table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    demand_para = doc.add_paragraph()
    demand_para.add_run('Demand Signal: ').bold = True
    demand_para.add_run(market.get('demandSignal', 'Unknown'))
    
    rationale = doc.add_paragraph()
    rationale.add_run('Rationale: ').bold = True
    rationale.add_run(market.get('demandRationale', ''))
    
    doc.add_page_break()
    
    # ============ FLAGS & QUESTIONS ============
    add_heading(doc, 'Diligence Flags')
    
    flags = fin.get('flags', [])
    if flags:
        table = doc.add_table(rows=len(flags) + 1, cols=4)
        table.style = 'Table Grid'
        
        for i, h in enumerate(['Item', 'Issue', 'Severity', 'Recommendation']):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[0].cells[i], "D6E3F8")
        
        for i, flag in enumerate(flags, 1):
            table.rows[i].cells[0].text = flag.get('item', '')
            table.rows[i].cells[1].text = flag.get('issue', '')
            table.rows[i].cells[2].text = flag.get('severity', '')
            table.rows[i].cells[3].text = flag.get('recommendation', '')
    else:
        doc.add_paragraph('No significant flags identified.')
    
    doc.add_paragraph()
    add_heading(doc, 'Questions for Seller/Broker')
    
    questions = fin.get('questions', [])
    if questions:
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q}")
    else:
        doc.add_paragraph('No additional questions at this time.')
    
    doc.add_paragraph()
    
    # ============ UNDERWRITING ASSUMPTIONS ============
    add_heading(doc, 'GGC Underwriting Assumptions')
    
    assumptions = [
        'Lot Rent Income: T3 annualized (T3 × 4)',
        'Other Income: T12 as reported',
        'Operating Expenses: T12 + 3% inflation adjustment',
        'Management Fee: 5% of Effective Gross Income',
        'CapEx Reserve: $50/unit annually',
        'Vacancy: 5% minimum (or actual if higher)'
    ]
    
    for assumption in assumptions:
        p = doc.add_paragraph(assumption)
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # ============ DISCLAIMER ============
    add_heading(doc, 'Disclaimer')
    
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        'This investment memo is prepared for internal use by Gary Group Capital and is based on '
        'information provided by the seller and publicly available data. All projections and analyses '
        'are estimates and should not be relied upon as guarantees of future performance. Investors '
        'should conduct their own due diligence before making any investment decisions. Gary Group '
        'Capital makes no representations or warranties regarding the accuracy or completeness of '
        'this information.'
    )
    disclaimer.paragraph_format.first_line_indent = Inches(0.5)
    for run in disclaimer.runs:
        run.font.size = Pt(10)
        run.font.italic = True
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def add_heading(doc, text):
    """Add a styled heading"""
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)
